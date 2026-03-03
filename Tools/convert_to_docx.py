#!/usr/bin/env python3
"""
Markdown to DOCX Converter
Converts chapter Markdown files to DOCX format for editing.
Preserves:
- ARMI system messages (as styled blockquotes)
- German special characters
- All formatting (headers, bold, italic)
"""

import argparse
import re
import sys
from pathlib import Path
from typing import Optional, List

try:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
except ImportError:
    print("Error: 'python-docx' package not found.")
    print("Install it with: pip install python-docx")
    sys.exit(1)


class DocxConverter:
    """Converts Markdown to DOCX format"""
    
    def __init__(self):
        pass
    
    def convert_file(self, input_path: Path, output_path: Optional[Path] = None) -> Path:
        """
        Convert a single Markdown file to DOCX
        
        Args:
            input_path: Path to input .md file
            output_path: Optional path to output .docx file
            
        Returns:
            Path to created DOCX file
        """
        if not input_path.exists():
            raise FileNotFoundError(f"Input file not found: {input_path}")
        
        # Read Markdown content
        content = input_path.read_text(encoding='utf-8')
        
        # Preprocess: Remove YAML frontmatter if present
        content = self._remove_frontmatter(content)
        
        # Preprocess: Remove word count footer
        content = self._remove_footer(content)
        
        # Create DOCX document
        doc = Document()
        
        # Setup styles
        self._setup_styles(doc)
        
        # Parse and add content
        self._parse_markdown_to_docx(content, doc)
        
        # Determine output path
        if output_path is None:
            output_dir = input_path.parent.parent / "Chapters-DOCX"
            output_dir.mkdir(exist_ok=True)
            output_path = output_dir / input_path.with_suffix('.docx').name
        
        # Save DOCX file
        doc.save(str(output_path))
        
        return output_path
    
    def _remove_frontmatter(self, content: str) -> str:
        """Remove YAML frontmatter if present"""
        if content.startswith('---'):
            # Find second --- and remove everything before it
            parts = content.split('---', 2)
            if len(parts) >= 3:
                return parts[2].lstrip()
        return content
    
    def _remove_footer(self, content: str) -> str:
        """Remove word count footer and trailing metadata"""
        # Remove lines like: **Word Count**: ~1,620
        content = re.sub(r'\*\*Word Count\*\*:.*', '', content)
        
        # Remove lines like: **Next Chapter**: [...]
        content = re.sub(r'\*\*Next Chapter\*\*:.*', '', content)
        
        # Remove multiple trailing newlines
        content = content.rstrip() + '\n'
        
        return content
    
    def _setup_styles(self, doc: Document):
        """Setup custom styles for the document"""
        styles = doc.styles
        
        # ARMI message style (for blockquotes)
        try:
            armi_style = styles.add_style('ARMI Message', WD_STYLE_TYPE.PARAGRAPH)
            armi_style.font.name = 'Courier New'
            armi_style.font.size = Pt(10)
            armi_style.font.color.rgb = RGBColor(44, 62, 80)
            armi_style.paragraph_format.left_indent = Inches(0.5)
            armi_style.paragraph_format.space_before = Pt(12)
            armi_style.paragraph_format.space_after = Pt(12)
        except ValueError:
            # Style already exists
            pass
    
    def _parse_markdown_to_docx(self, content: str, doc: Document):
        """Parse Markdown content and add to DOCX document"""
        lines = content.split('\n')
        i = 0
        in_blockquote = False
        blockquote_lines = []
        
        while i < len(lines):
            line = lines[i]
            
            # Handle blockquotes (ARMI messages)
            if line.startswith('>'):
                in_blockquote = True
                # Remove > prefix
                blockquote_text = line.lstrip('>').strip()
                blockquote_lines.append(blockquote_text)
                i += 1
                continue
            elif in_blockquote:
                # End of blockquote
                self._add_blockquote(doc, blockquote_lines)
                blockquote_lines = []
                in_blockquote = False
            
            # Handle headers
            if line.startswith('# '):
                doc.add_heading(line[2:], level=1)
            elif line.startswith('## '):
                doc.add_heading(line[3:], level=2)
            elif line.startswith('### '):
                doc.add_heading(line[4:], level=3)
            elif line.startswith('---') or line.startswith('***'):
                # Horizontal rule - add a simple line break or styled separator
                p = doc.add_paragraph()
                p.add_run('─' * 30).font.color.rgb = RGBColor(150, 150, 150)
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif line.strip():
                # Regular paragraph - handle inline formatting
                self._add_formatted_paragraph(doc, line)
            elif not line.strip():
                # Empty line - skip but allow paragraph breaks naturally
                pass
            
            i += 1
        
        # Handle any remaining blockquote
        if blockquote_lines:
            self._add_blockquote(doc, blockquote_lines)
    
    def _add_blockquote(self, doc: Document, lines: List[str]):
        """Add a blockquote (ARMI message) to the document"""
        text = '\n'.join(lines)
        try:
            p = doc.add_paragraph(style='ARMI Message')
        except KeyError:
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.5)
        
        # Apply formatting within the blockquote
        self._add_formatted_runs(p, text)
    
    def _add_formatted_paragraph(self, doc: Document, text: str):
        """Add a paragraph with inline formatting"""
        p = doc.add_paragraph()
        self._add_formatted_runs(p, text)
    
    def _add_formatted_runs(self, paragraph, text: str):
        """Add formatted runs to a paragraph (bold, italic)"""
        # Pattern for bold and italic
        # ***text*** or ___text___ = bold+italic
        # **text** or __text__ = bold
        # *text* or _text_ = italic
        
        pattern = r'(\*\*\*(.+?)\*\*\*|\*\*(.+?)\*\*|\*(.+?)\*)'
        
        last_end = 0
        for match in re.finditer(pattern, text):
            # Add text before match
            if match.start() > last_end:
                paragraph.add_run(text[last_end:match.start()])
            
            # Add formatted text
            if match.group(2):  # Bold+Italic (***text***)
                run = paragraph.add_run(match.group(2))
                run.bold = True
                run.italic = True
            elif match.group(3):  # Bold (**text**)
                run = paragraph.add_run(match.group(3))
                run.bold = True
            elif match.group(4):  # Italic (*text*)
                run = paragraph.add_run(match.group(4))
                run.italic = True
            
            last_end = match.end()
        
        # Add remaining text
        if last_end < len(text):
            paragraph.add_run(text[last_end:])


def main():
    """Main CLI entry point"""
    parser = argparse.ArgumentParser(
        description='Convert Markdown chapters to DOCX format'
    )
    parser.add_argument(
        'input',
        nargs='?',
        help='Input Markdown file or directory'
    )
    parser.add_argument(
        '--range',
        '-r',
        help='Chapter range to convert (e.g., "11-50")'
    )
    parser.add_argument(
        '--all',
        action='store_true',
        help='Convert all chapters in Story/Chapters/'
    )
    parser.add_argument(
        '--output',
        '-o',
        help='Output directory (default: Story/Chapters-DOCX/)'
    )
    
    args = parser.parse_args()
    
    # Determine project root (where Tools/ directory is located)
    script_dir = Path(__file__).parent
    project_root = script_dir.parent
    
    converter = DocxConverter()
    
    if args.all or args.range:
        # Convert chapters
        chapters_dir = project_root / "Story" / "Chapters"
        if not chapters_dir.exists():
            print(f"Error: Chapters directory not found: {chapters_dir}")
            sys.exit(1)
        
        markdown_files = sorted(chapters_dir.glob("Chapter_*.md"))
        
        if not markdown_files:
            print(f"No chapter files found in {chapters_dir}")
            sys.exit(1)
        
        # Filter by range if specified
        if args.range:
            try:
                start, end = map(int, args.range.split('-'))
                filtered_files = []
                for md_file in markdown_files:
                    # Extract chapter number from filename
                    match = re.search(r'Chapter_(\d+)', md_file.name)
                    if match:
                        chapter_num = int(match.group(1))
                        if start <= chapter_num <= end:
                            filtered_files.append(md_file)
                markdown_files = sorted(filtered_files, key=lambda f: int(re.search(r'Chapter_(\d+)', f.name).group(1)))
            except ValueError:
                print(f"Error: Invalid range format. Use 'start-end' (e.g., '11-50')")
                sys.exit(1)
        
        print(f"Found {len(markdown_files)} chapter files to convert")
        converted_count = 0
        
        for md_file in markdown_files:
            try:
                output_path = converter.convert_file(md_file)
                print(f"[OK] {md_file.name} -> {output_path.name}")
                converted_count += 1
            except Exception as e:
                print(f"[FAIL] {md_file.name}: {e}")
        
        print(f"\nConverted {converted_count}/{len(markdown_files)} files")
        
    elif args.input:
        # Convert single file
        input_path = Path(args.input)
        
        if not input_path.is_absolute():
            # Try relative to current directory first
            if input_path.exists():
                pass
            # Try relative to project root
            elif (project_root / input_path).exists():
                input_path = project_root / input_path
            else:
                print(f"Error: File not found: {input_path}")
                sys.exit(1)
        
        try:
            output_path = converter.convert_file(input_path)
            print(f"[OK] Converted: {output_path}")
        except Exception as e:
            print(f"[FAIL] Error: {e}")
            sys.exit(1)
    
    else:
        parser.print_help()
        sys.exit(1)


if __name__ == '__main__':
    main()
